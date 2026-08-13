"""
VERIFICATION DES RAPPORTS EXPORTES — les quatre langues
PrevuFlow — SMD Global Consulting LLC

Le rapport est ce qu'un dirigeant transmet a son banquier : c'est la
vitrine du produit hors de l'application. Il est reste francais jusqu'au
13 aout 2026, alors meme que l'ecran, lui, etait traduit.

Ce test ouvre vraiment les fichiers produits — classeur Excel, document
Word — et lit ce qu'ils contiennent. Verifier que la fonction ne leve pas
d'exception ne prouverait rien : un rapport peut se generer parfaitement
et rester en francais.

Lancer :  py test_export_langues.py
"""

from __future__ import annotations

import io
import os
import sys
from datetime import date, timedelta
from decimal import Decimal

os.environ["PREVUFLOW_HEBERGE"] = "1"

import langues as lg                                            # noqa: E402
from analyse_entreprise import Facture, calculer_indicateurs    # noqa: E402
from analyse_lisible import analyser_lisible                    # noqa: E402
from import_intelligent import Mouvement                        # noqa: E402

import export_rapport as ex                                     # noqa: E402

echecs = 0


def verifier(titre: str, condition: bool) -> None:
    global echecs
    if condition:
        print(f"  ok    {titre}")
    else:
        print(f"  ECHEC {titre}")
        echecs += 1


def traducteur(code: str):
    return (lambda cle, **v: lg.traduire(cle, code, **v),
            lambda valeur: lg.formater_montant(valeur, "EUR", code))


JOUR = date.today()
MOUVEMENTS = (
    [Mouvement(JOUR - timedelta(days=n * 7), "Groceries", Decimal("-62.40"))
     for n in range(8)]
    + [Mouvement(JOUR - timedelta(days=n * 30), "Salary", Decimal("2800"))
       for n in range(3)])
SYNTHESE = analyser_lisible(MOUVEMENTS)
INDICATEURS = calculer_indicateurs(
    mouvements=MOUVEMENTS,
    factures_clients=[Facture(date_emission=JOUR - timedelta(days=60),
                              tiers="Alpha", montant=Decimal("9000"),
                              sens="client",
                              echeance=JOUR - timedelta(days=30))],
    tresorerie=Decimal("9000"))

# Des phrases assez longues pour ne pas se confondre d'une langue a l'autre.
FR = {v for v in lg._charger("fr").values() if len(v) > 16 and "{" not in v}


def textes_du_classeur(octets: bytes) -> tuple[list[str], list[str]]:
    from openpyxl import load_workbook

    classeur = load_workbook(io.BytesIO(octets))
    valeurs: list[str] = []
    for feuille in classeur.worksheets:
        for ligne in feuille.iter_rows(values_only=True):
            valeurs += [str(v) for v in ligne if isinstance(v, str)]
    return classeur.sheetnames, valeurs


def textes_du_word(octets: bytes) -> list[str]:
    from docx import Document

    doc = Document(io.BytesIO(octets))
    valeurs = [p.text for p in doc.paragraphs if p.text.strip()]
    for tableau in doc.tables:
        for ligne in tableau.rows:
            valeurs += [c.text for c in ligne.cells if c.text.strip()]
    return valeurs


print("1. Le classeur Excel du profil Particulier")

for code in ("fr", "en", "es", "zh"):
    t, nombre = traducteur(code)
    feuilles, textes = textes_du_classeur(
        ex.exporter_excel(SYNTHESE, MOUVEMENTS, devise="EUR", t=t,
                          nombre=nombre))
    verifier(f"{code} : le classeur a ses quatre feuilles", len(feuilles) == 4)
    verifier(f"{code} : les onglets sont traduits ({feuilles})",
             feuilles[0] == lg.traduire("exp.synthese", code))
    if code == "fr":
        continue
    restes = sorted({p for p in FR if any(p in v for v in textes)})
    verifier(f"{code} : aucune phrase francaise dans le classeur", not restes)
    for r in restes[:5]:
        print(f"          · {r[:70]}")


print("\n2. Le document Word")

for code in ("fr", "en", "es", "zh"):
    t, nombre = traducteur(code)
    textes = textes_du_word(
        ex.exporter_word(SYNTHESE, devise="EUR", t=t, nombre=nombre))
    verifier(f"{code} : le document a du contenu", len(textes) > 10)
    verifier(f"{code} : le titre est traduit",
             any(lg.traduire("exp.titre_defaut", code) in v for v in textes))
    if code == "fr":
        continue
    restes = sorted({p for p in FR if any(p in v for v in textes)})
    verifier(f"{code} : aucune phrase francaise dans le Word", not restes)
    for r in restes[:5]:
        print(f"          · {r[:70]}")


print("\n3. Le classeur du profil Entreprise")

for code in ("fr", "en", "es", "zh"):
    t, nombre = traducteur(code)
    feuilles, textes = textes_du_classeur(
        ex.exporter_entreprise_excel(INDICATEURS, SYNTHESE, MOUVEMENTS,
                                     devise="EUR", t=t, nombre=nombre))
    verifier(f"{code} : les blocs d'indicateurs sont traduits",
             any(lg.traduire("exp.bloc_activite", code) in v for v in textes))
    if code == "fr":
        continue
    restes = sorted({p for p in FR if any(p in v for v in textes)})
    verifier(f"{code} : aucune phrase francaise dans le classeur Entreprise",
             not restes)
    for r in restes[:5]:
        print(f"          · {r[:70]}")


print("\n4. Le PDF se produit dans les quatre langues")

# Le PDF n'est pas relisible sans dependance supplementaire : on verifie
# qu'il se construit et qu'il pese quelque chose. Son contenu vient des
# memes cles que le Word, deja controle ci-dessus.
for code in ("fr", "en", "es", "zh"):
    t, nombre = traducteur(code)
    try:
        octets = ex.exporter_pdf(SYNTHESE, devise="EUR", t=t, nombre=nombre)
        verifier(f"{code} : PDF produit ({len(octets)} octets)",
                 len(octets) > 1500)
    except Exception as erreur:
        verifier(f"{code} : PDF produit ({type(erreur).__name__} {erreur})",
                 False)


print("\n5. Le montant ne porte pas deux fois sa devise")

# Defaut vu a la relecture : le formateur rend « 4 008,34 € » et le texte
# ajoutait le symbole une seconde fois.
t, nombre = traducteur("fr")
textes = textes_du_word(ex.exporter_word(SYNTHESE, devise="EUR", t=t,
                                         nombre=nombre))
verifier("aucun « € € » dans le rapport",
         not any("€ €" in v for v in textes))
verifier("aucun double espace dans les phrases",
         not any("  " in v for v in textes))


print("\n" + "=" * 62)
print("Toutes les verifications sont passees."
      if not echecs else f"{echecs} echec(s).")
sys.exit(1 if echecs else 0)
