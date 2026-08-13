"""
EXPORT DU RAPPORT FINANCIER
PrevuFlow — SMD Global Consulting LLC

Produit un rapport telechargeable en Excel ou en PDF a partir de l'analyse.

Dependances : openpyxl (Excel) · reportlab (PDF)
"""

from __future__ import annotations

import io
from datetime import date, datetime
from decimal import Decimal

def _texte(cle: str, **variables) -> str:
    """
    Repli quand l'appelant ne fournit pas sa propre traduction.

    Ce module fabrique des fichiers ; il ne connait pas la langue de son
    lecteur. Sans traduction fournie, il repond en francais.
    """
    from langues import traduire
    return traduire(cle, "fr", **variables)


BLEU = "1F4E79"
VERT = "1F7244"
ROUGE = "C0392B"
GRIS = "F2F2F2"
POLICE = "Arial"


def _phrases(source, t, montant) -> list[tuple[str, str]]:
    """
    Les constats du moteur, mis en forme pour un document.

    Le formateur de montants porte deja la devise — « 4 008,34 € ». Les
    textes, eux, prevoient un symbole separe pour l'ecran. On passe donc un
    symbole vide et on referme l'espace laisse : sans cela le rapport
    afficherait « 4 008,34 € € ».
    """
    return [(niveau, " ".join(texte.split()))
            for niveau, texte in source.messages(t, montant, "")]


def _symbole(devise: str = "EUR") -> str:
    return {"EUR": "€", "USD": "$", "GBP": "£", "XOF": "FCFA",
            "CAD": "C$", "CHF": "CHF", "MAD": "DH"}.get(devise, devise)


def _euro(v, devise: str = "EUR") -> str:
    symbole = _symbole(devise)
    texte = f"{abs(float(v)):,.2f}".replace(",", " ").replace(".", ",")
    return f"{'-' if float(v) < 0 else ''}{texte} {symbole}"


# ---------------------------------------------------------------------------
# EXCEL
# ---------------------------------------------------------------------------

def exporter_excel(synthese, mouvements, projection=None,
                   devise: str = "EUR", titre: str = "", t=None,
                   nombre=None) -> bytes:
    """Renvoie un classeur Excel : synthèse, postes, catégories, opérations."""
    t = t or _texte
    titre = titre or t("exp.titre_defaut")
    montant = nombre or (lambda v: _euro(v, devise))
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter

    wb = Workbook()
    bord = Border(*[Side(style="thin", color="D9D9D9")] * 4)

    def entete(ws, colonnes, ligne=1):
        for i, h in enumerate(colonnes, 1):
            c = ws.cell(row=ligne, column=i, value=h)
            c.font = Font(name=POLICE, size=10, bold=True, color="FFFFFF")
            c.fill = PatternFill("solid", fgColor=BLEU)
            c.alignment = Alignment(horizontal="center", vertical="center",
                                    wrap_text=True)
        ws.row_dimensions[ligne].height = 26

    # --- Feuille 1 : Synthèse ---------------------------------------------
    ws = wb.active
    ws.title = t("exp.synthese")
    ws["A1"] = titre
    ws["A1"].font = Font(name=POLICE, size=15, bold=True, color=BLEU)
    ws["A2"] = t("exp.periode", debut=synthese.debut.strftime("%d/%m/%Y"),
                 fin=synthese.fin.strftime("%d/%m/%Y"),
                 n=synthese.nb_operations, mois=synthese.nb_mois)
    ws["A2"].font = Font(name=POLICE, size=9, italic=True, color="666666")

    lignes = [
        ("", ""),
        (t("exp.chiffres_cles"), t("exp.par_mois")),
        (t("exp.rentre"), float(synthese.entrees_par_mois)),
        (t("exp.sort"), float(synthese.sorties_par_mois)),
        (t("exp.reste"), float(synthese.reste_par_mois)),
        ("", ""),
        (t("exp.dont_fixes"), float(synthese.charges_fixes)),
        (t("exp.dont_variables"), float(synthese.depenses_variables)),
        ("", ""),
        (t("exp.sur_periode"), ""),
        (t("exp.total_encaisse"), float(synthese.entrees)),
        (t("exp.total_depense"), float(synthese.sorties)),
        (t("exp.resultat_periode"), float(synthese.solde_periode)),
        ("", ""),
        (t("exp.indicateurs"), ""),
        (t("exp.part_epargne"), f"{synthese.taux_epargne:.0f} %"),
        (t("exp.part_fixe"), f"{synthese.part_fixe:.0f} %"),
    ]
    for i, (lib, val) in enumerate(lignes, 4):
        cl = ws.cell(row=i, column=1, value=lib)
        titre_section = lib.isupper() and lib
        en_relief = lib in (t("exp.reste"), t("exp.resultat_periode"))
        cl.font = Font(name=POLICE, size=11,
                       bold=bool(titre_section) or en_relief,
                       color=BLEU if titre_section else "000000")
        cv = ws.cell(row=i, column=2, value=val)
        cv.font = Font(name=POLICE, size=11, bold=en_relief,
                       color=(ROUGE if isinstance(val, float) and val < 0 else "000000"))
        if isinstance(val, float):
            cv.number_format = '#,##0.00 "€";[Red]-#,##0.00 "€"'
        cv.alignment = Alignment(horizontal="right")

    depart = len(lignes) + 5
    ws.cell(row=depart, column=1, value=t("exp.ce_que_ca_veut")).font = \
        Font(name=POLICE, size=11, bold=True, color=BLEU)
    for i, (niveau, texte) in enumerate(_phrases(synthese, t, montant), depart + 1):
        couleur = {"alerte": ROUGE, "attention": "B7791F",
                   "bon": VERT}.get(niveau, "000000")
        c = ws.cell(row=i, column=1, value=texte)
        c.font = Font(name=POLICE, size=10, color=couleur)
        ws.merge_cells(start_row=i, start_column=1, end_row=i, end_column=4)
    ws.column_dimensions["A"].width = 46
    ws.column_dimensions["B"].width = 18

    # --- Feuille 2 : Postes ------------------------------------------------
    ws2 = wb.create_sheet(t("exp.postes"))
    entete(ws2, [t("exp.poste"), t("col.categorie"), t("col.type"),
                 t("col.nombre"), t("exp.total_periode"), t("exp.moyenne"),
                 t("col.par_mois")])
    for i, p in enumerate(synthese.postes, 2):
        for j, v in enumerate([
                p.nom, p.categorie,
                (t("col.charge_fixe") if p.fixe
                 else t("app.entree") if p.est_une_entree
                 else t("col.variable")),
                p.nombre, float(p.total), float(p.moyenne), float(p.par_mois)], 1):
            c = ws2.cell(row=i, column=j, value=v)
            c.font = Font(name=POLICE, size=10)
            c.border = bord
            if j >= 5:
                c.number_format = '#,##0.00;[Red]-#,##0.00'
            if i % 2 == 0:
                c.fill = PatternFill("solid", fgColor=GRIS)
    for i, w in enumerate([34, 26, 13, 9, 15, 13, 13], 1):
        ws2.column_dimensions[get_column_letter(i)].width = w
    ws2.freeze_panes = "A2"
    ws2.auto_filter.ref = f"A1:G{len(synthese.postes) + 1}"

    # --- Feuille 3 : Catégories -------------------------------------------
    ws3 = wb.create_sheet(t("exp.categories"))
    entete(ws3, [t("col.categorie"), t("exp.total_periode"), t("col.par_mois"),
                 t("exp.part_depenses")])
    total_sorties = abs(float(synthese.sorties)) or 1
    for i, (cat, tot) in enumerate(synthese.par_categorie.items(), 2):
        part = abs(float(tot)) / total_sorties if float(tot) < 0 else 0
        for j, v in enumerate([cat, float(tot),
                               float(tot) / max(synthese.nb_mois, 0.5), part], 1):
            c = ws3.cell(row=i, column=j, value=v)
            c.font = Font(name=POLICE, size=10)
            c.border = bord
            if j in (2, 3):
                c.number_format = '#,##0.00;[Red]-#,##0.00'
            if j == 4:
                c.number_format = "0.0 %"
    for i, w in enumerate([32, 16, 14, 18], 1):
        ws3.column_dimensions[get_column_letter(i)].width = w
    ws3.freeze_panes = "A2"

    # --- Feuille 4 : Opérations -------------------------------------------
    ws4 = wb.create_sheet(t("exp.operations"))
    entete(ws4, [t("col.date"), t("exp.libelle"), t("col.categorie"),
                 t("col.montant")])
    for i, m in enumerate(sorted(mouvements, key=lambda x: x.jour), 2):
        from analyse_lisible import categoriser
        for j, v in enumerate([m.jour, m.libelle,
                               categoriser(m.libelle, m.montant),
                               float(m.montant)], 1):
            c = ws4.cell(row=i, column=j, value=v)
            c.font = Font(name=POLICE, size=10)
            c.border = bord
            if j == 1:
                c.number_format = "DD/MM/YYYY"
            if j == 4:
                c.number_format = '#,##0.00;[Red]-#,##0.00'
    for i, w in enumerate([13, 52, 26, 14], 1):
        ws4.column_dimensions[get_column_letter(i)].width = w
    ws4.freeze_panes = "A2"
    ws4.auto_filter.ref = f"A1:D{len(mouvements) + 1}"

    # --- Feuille 5 : Prévision (facultative) ------------------------------
    if projection:
        ws5 = wb.create_sheet(t("exp.prevision"))
        entete(ws5, [t("col.date"), t("exp.entrees"), t("exp.sorties"),
                     t("exp.solde_projete"), t("exp.operations")])
        for i, j_ in enumerate(projection, 2):
            for k, v in enumerate([j_.jour, float(j_.entrees), float(j_.sorties),
                                   float(j_.solde), " · ".join(j_.operations)], 1):
                c = ws5.cell(row=i, column=k, value=v)
                c.font = Font(name=POLICE, size=10)
                c.border = bord
                if k == 1:
                    c.number_format = "DD/MM/YYYY"
                if k in (2, 3, 4):
                    c.number_format = '#,##0.00;[Red]-#,##0.00'
                if k == 4 and float(j_.solde) < 0:
                    c.font = Font(name=POLICE, size=10, bold=True, color=ROUGE)
        for i, w in enumerate([13, 14, 14, 16, 46], 1):
            ws5.column_dimensions[get_column_letter(i)].width = w
        ws5.freeze_panes = "A2"

    flux = io.BytesIO()
    wb.save(flux)
    return flux.getvalue()


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------

def exporter_pdf(synthese, devise: str = "EUR", titre: str = "",
                 auteur: str = "SMD Global Consulting LLC", t=None,
                 nombre=None) -> bytes:
    """Renvoie un rapport PDF d'une à deux pages, lisible par un non-financier."""
    t = t or _texte
    titre = titre or t("exp.titre_defaut")
    montant = nombre or (lambda v: _euro(v, devise))
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.platypus import (Paragraph, SimpleDocTemplate, Spacer, Table,
                                    TableStyle)

    flux = io.BytesIO()
    doc = SimpleDocTemplate(flux, pagesize=A4,
                            leftMargin=18 * mm, rightMargin=18 * mm,
                            topMargin=16 * mm, bottomMargin=16 * mm,
                            title=titre, author=auteur)
    styles = getSampleStyleSheet()
    st_titre = ParagraphStyle("t", parent=styles["Title"], fontName="Helvetica-Bold",
                              fontSize=19, textColor=colors.HexColor("#" + BLEU),
                              alignment=0, spaceAfter=2)
    st_sous = ParagraphStyle("s", parent=styles["Normal"], fontSize=9,
                             textColor=colors.HexColor("#666666"), spaceAfter=14)
    st_h = ParagraphStyle("h", parent=styles["Heading2"], fontName="Helvetica-Bold",
                          fontSize=12, textColor=colors.HexColor("#" + BLEU),
                          spaceBefore=14, spaceAfter=6)
    st_n = ParagraphStyle("n", parent=styles["Normal"], fontSize=10, leading=15)

    contenu = [
        Paragraph(titre, st_titre),
        Paragraph(t("exp.periode",
                    debut=synthese.debut.strftime("%d/%m/%Y"),
                    fin=synthese.fin.strftime("%d/%m/%Y"),
                    n=synthese.nb_operations, mois=synthese.nb_mois), st_sous),
    ]

    # Chiffres clés
    contenu.append(Paragraph(t("exp.chiffres_par_mois"), st_h))
    donnees = [
        [t("exp.rentre"), montant(synthese.entrees_par_mois)],
        [t("exp.sort"), montant(synthese.sorties_par_mois)],
        [t("exp.reste"), montant(synthese.reste_par_mois)],
        [t("exp.dont_fixes"), montant(synthese.charges_fixes)],
        [t("exp.dont_variables"), montant(synthese.depenses_variables)],
    ]
    tableau = Table(donnees, colWidths=[105 * mm, 60 * mm])
    tableau.setStyle(TableStyle([
        ("FONT", (0, 0), (-1, -1), "Helvetica", 10),
        ("FONT", (0, 2), (-1, 2), "Helvetica-Bold", 11),
        ("TEXTCOLOR", (1, 2), (1, 2),
         colors.HexColor("#" + (ROUGE if synthese.reste_par_mois < 0 else VERT))),
        ("ALIGN", (1, 0), (1, -1), "RIGHT"),
        ("LINEBELOW", (0, 0), (-1, -2), 0.4, colors.HexColor("#DDDDDD")),
        ("LINEBELOW", (0, 2), (-1, 2), 1.0, colors.HexColor("#" + BLEU)),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("LEFTINDENT", (0, 3), (0, 4), 12),
    ]))
    contenu += [tableau, Spacer(1, 4)]

    # Ce que ça veut dire
    contenu.append(Paragraph(t("exp.ce_que_ca_veut"), st_h))
    for niveau, texte in _phrases(synthese, t, montant):
        puce = {"alerte": "▲", "attention": "▲", "bon": "●"}.get(niveau, "•")
        couleur = {"alerte": ROUGE, "attention": "B7791F", "bon": VERT}.get(niveau, "333333")
        contenu.append(Paragraph(
            f'<font color="#{couleur}">{puce}</font> {texte}', st_n))
        contenu.append(Spacer(1, 3))

    # Principaux postes
    contenu.append(Paragraph(t("exp.principaux_postes"), st_h))
    lignes = [[t("exp.poste"), t("col.categorie"), t("col.par_mois")]]
    for p in synthese.postes[:12]:
        lignes.append([p.nom[:30], p.categorie[:24], montant(p.par_mois)])
    t2 = Table(lignes, colWidths=[62 * mm, 62 * mm, 41 * mm], repeatRows=1)
    t2.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#" + BLEU)),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONT", (0, 0), (-1, 0), "Helvetica-Bold", 9),
        ("FONT", (0, 1), (-1, -1), "Helvetica", 9),
        ("ALIGN", (2, 0), (2, -1), "RIGHT"),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F7F7F7")]),
        ("GRID", (0, 0), (-1, -1), 0.3, colors.HexColor("#DDDDDD")),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
    ]))
    contenu.append(t2)

    # Répartition par catégorie
    contenu.append(Paragraph(t("exp.repartition_categorie"), st_h))
    total_sorties = abs(float(synthese.sorties)) or 1
    lignes3 = [[t("col.categorie"), t("exp.total_periode"), t("exp.part")]]
    for cat, tot in list(synthese.par_categorie.items())[:12]:
        part = f"{abs(float(tot)) / total_sorties * 100:.0f} %" if float(tot) < 0 else "—"
        lignes3.append([cat[:34], montant(tot), part])
    t3 = Table(lignes3, colWidths=[85 * mm, 50 * mm, 30 * mm], repeatRows=1)
    t3.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#" + BLEU)),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONT", (0, 0), (-1, 0), "Helvetica-Bold", 9),
        ("FONT", (0, 1), (-1, -1), "Helvetica", 9),
        ("ALIGN", (1, 0), (-1, -1), "RIGHT"),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F7F7F7")]),
        ("GRID", (0, 0), (-1, -1), 0.3, colors.HexColor("#DDDDDD")),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
    ]))
    contenu.append(t3)

    contenu += [
        Spacer(1, 12),
        Paragraph('<font size="7.5" color="#888888">'
                  + t("exp.pied_pdf",
                      date=datetime.now().strftime("%d/%m/%Y"), auteur=auteur)
                  + "</font>", st_n),
    ]

    doc.build(contenu)
    return flux.getvalue()


# ---------------------------------------------------------------------------
# WORD — rapport modifiable
# ---------------------------------------------------------------------------

def exporter_word(synthese, devise: str = "EUR", titre: str = "",
                  auteur: str = "SMD Global Consulting LLC", t=None,
                  nombre=None) -> bytes:
    """
    Rapport Word modifiable, pour que l'utilisateur puisse le personnaliser
    avant de le transmettre a son banquier, son comptable ou son associe.

    Dependance : python-docx
    """
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor, Cm

    t = t or _texte
    titre = titre or t("exp.titre_defaut")
    montant = nombre or (lambda v: _euro(v, devise))

    def _rgb(hexa: str) -> RGBColor:
        return RGBColor(int(hexa[0:2], 16), int(hexa[2:4], 16), int(hexa[4:6], 16))

    doc = Document()

    for section in doc.sections:
        section.top_margin = section.bottom_margin = Cm(2)
        section.left_margin = section.right_margin = Cm(2)

    normal = doc.styles["Normal"]
    normal.font.name = POLICE
    normal.font.size = Pt(10.5)

    # --- Titre ---
    para = doc.add_paragraph()
    r = para.add_run(titre)
    r.font.size = Pt(20)
    r.font.bold = True
    r.font.color.rgb = _rgb(BLEU)

    st_ = doc.add_paragraph()
    r = st_.add_run(t("exp.etabli_le",
                      date=date.today().strftime("%d/%m/%Y"), auteur=auteur))
    r.font.size = Pt(9)
    r.font.color.rgb = _rgb("808080")

    # --- Chiffres cles ---
    doc.add_paragraph()
    h = doc.add_paragraph()
    r = h.add_run(t("exp.chiffres_chaque_mois"))
    r.font.size = Pt(13); r.font.bold = True; r.font.color.rgb = _rgb(BLEU)

    lignes = [
        (t("exp.rentre"), synthese.entrees_par_mois, VERT, False),
        (t("exp.sort"), -abs(synthese.sorties_par_mois), ROUGE, False),
        (t("exp.reste"), synthese.reste_par_mois,
         VERT if float(synthese.reste_par_mois) >= 0 else ROUGE, False),
        (t("exp.dont_fixes"), -abs(synthese.charges_fixes), "404040", True),
        (t("exp.dont_variables"), -abs(synthese.depenses_variables),
         "404040", True),
    ]

    tab = doc.add_table(rows=0, cols=2)
    tab.style = "Table Grid"
    for libelle, valeur, couleur, en_retrait in lignes:
        row = tab.add_row().cells
        p = row[0].paragraphs[0]
        rr = p.add_run(libelle)
        rr.font.size = Pt(10.5)
        rr.font.italic = en_retrait
        p2 = row[1].paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        rr2 = p2.add_run(montant(valeur))
        rr2.font.size = Pt(10.5); rr2.font.bold = True
        rr2.font.color.rgb = _rgb(couleur)

    # --- Ce que ca veut dire ---
    doc.add_paragraph()
    h = doc.add_paragraph()
    r = h.add_run(t("exp.ce_que_cela_veut"))
    r.font.size = Pt(13); r.font.bold = True; r.font.color.rgb = _rgb(BLEU)

    for niveau, texte in _phrases(synthese, t, montant):
        p = doc.add_paragraph(style="List Bullet")
        rr = p.add_run(texte)
        rr.font.size = Pt(10.5)
        if niveau == "alerte":
            rr.font.color.rgb = _rgb(ROUGE); rr.font.bold = True
        elif niveau == "bon":
            rr.font.color.rgb = _rgb(VERT)

    # --- Principaux postes ---
    doc.add_paragraph()
    h = doc.add_paragraph()
    r = h.add_run(t("exp.principaux_postes"))
    r.font.size = Pt(13); r.font.bold = True; r.font.color.rgb = _rgb(BLEU)

    tp = doc.add_table(rows=1, cols=4)
    tp.style = "Table Grid"
    for i, entete in enumerate([t("exp.poste"), t("col.categorie"),
                                t("exp.nature"), t("col.par_mois")]):
        c = tp.rows[0].cells[i].paragraphs[0]
        if i == 3:
            c.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        rr = c.add_run(entete)
        rr.font.bold = True; rr.font.size = Pt(10)
        rr.font.color.rgb = _rgb("FFFFFF")
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), BLEU)
        tp.rows[0].cells[i]._tc.get_or_add_tcPr().append(shd)

    for p_ in synthese.postes[:15]:
        row = tp.add_row().cells
        nature = (t("app.entree") if p_.est_une_entree
                  else t("col.charge_fixe") if p_.fixe else t("col.variable"))
        for i, val in enumerate([p_.nom, p_.categorie, nature,
                                 montant(p_.par_mois)]):
            par = row[i].paragraphs[0]
            if i == 3:
                par.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            rr = par.add_run(str(val))
            rr.font.size = Pt(9.5)
            if i == 3:
                rr.font.bold = True
                rr.font.color.rgb = _rgb(VERT if p_.est_une_entree else ROUGE)

    # --- Repartition par categorie ---
    doc.add_paragraph()
    h = doc.add_paragraph()
    r = h.add_run(t("exp.repartition_depenses"))
    r.font.size = Pt(13); r.font.bold = True; r.font.color.rgb = _rgb(BLEU)

    depenses = sorted(((k, abs(float(v))) for k, v in synthese.par_categorie.items()
                       if float(v) < 0), key=lambda x: -x[1])
    total = sum(v for _, v in depenses) or 1

    tc = doc.add_table(rows=0, cols=3)
    tc.style = "Table Grid"
    for nom, valeur in depenses:
        row = tc.add_row().cells
        row[0].paragraphs[0].add_run(nom).font.size = Pt(9.5)
        p2 = row[1].paragraphs[0]; p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p2.add_run(montant(-valeur)).font.size = Pt(9.5)
        p3 = row[2].paragraphs[0]; p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p3.add_run(f"{valeur / total:.0%}").font.size = Pt(9.5)

    # --- Pied ---
    doc.add_paragraph()
    p = doc.add_paragraph()
    rr = p.add_run(t("exp.pied_word"))
    rr.font.size = Pt(8); rr.font.italic = True
    rr.font.color.rgb = _rgb("808080")

    tampon = io.BytesIO()
    doc.save(tampon)
    return tampon.getvalue()


# ---------------------------------------------------------------------------
# EXCEL — profil entreprise
# ---------------------------------------------------------------------------

def exporter_entreprise_excel(indicateurs, synthese=None, mouvements=None,
                              devise: str = "EUR", titre: str = "",
                              t=None, nombre=None) -> bytes:
    """
    Classeur destine a un dirigeant, un banquier ou un investisseur.

    Feuille 1 : les indicateurs et leur lecture en clair.
    Feuille 2 : la repartition du chiffre d'affaires par client.
    Feuille 3 : les postes de depense, si un releve a ete importe.
    Feuille 4 : le detail des mouvements.
    """
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
    from openpyxl.utils import get_column_letter

    t = t or _texte
    titre = titre or t("exp.titre_defaut")
    montant = nombre or (lambda v: _euro(v, devise))

    i = indicateurs
    classeur = Workbook()

    titre_f = Font(name=POLICE, size=15, bold=True, color=BLEU)
    section_f = Font(name=POLICE, size=11, bold=True, color="FFFFFF")
    normal_f = Font(name=POLICE, size=10)
    gras_f = Font(name=POLICE, size=10, bold=True)
    fond_bleu = PatternFill("solid", fgColor=BLEU)
    fond_gris = PatternFill("solid", fgColor=GRIS)
    bord = Border(*(Side(style="thin", color="D9D9D9"),) * 4)
    droite = Alignment(horizontal="right")
    fmt = f'# ##0.00 "{_symbole(devise)}"'

    def entete(f, ligne, colonnes, largeurs):
        for n, (texte, larg) in enumerate(zip(colonnes, largeurs), start=1):
            c = f.cell(row=ligne, column=n, value=texte)
            c.font = section_f
            c.fill = fond_bleu
            f.column_dimensions[get_column_letter(n)].width = larg

    # ---------------- Feuille 1 : indicateurs ----------------
    f = classeur.active
    f.title = t("exp.indicateurs_feuille")
    f["A1"] = titre
    f["A1"].font = titre_f
    f["A2"] = t("exp.periode_ent", debut=i.debut.strftime("%d/%m/%Y"),
                fin=i.fin.strftime("%d/%m/%Y"), mois=f"{i.nb_mois:.1f}")
    f["A2"].font = Font(name=POLICE, size=9, italic=True, color="808080")
    f["A3"] = t("exp.etabli_simple", date=date.today().strftime("%d/%m/%Y"))
    f["A3"].font = Font(name=POLICE, size=9, color="808080")

    ligne = 5
    blocs = [
        (t("exp.bloc_activite"), [
            (t("ent.encaisse"), i.encaissements_par_mois,
             t("ent.note_encaisse")),
            (t("ent.decaisse"), i.decaissements_par_mois,
             t("ent.note_decaisse")),
            (t("ent.resultat"), i.resultat_par_mois,
             t("ent.note_resultat", p=f"{i.marge:.0f}")),
            (t("exp.croissance"),
             f"{i.croissance:.1f} %" if i.croissance is not None else "—",
             t("exp.note_croissance")),
        ]),
        (t("exp.bloc_couts"), [
            (t("ent.charges_fixes"), i.charges_fixes,
             t("exp.note_part_fixe", p=f"{i.part_fixe:.0f}")),
            (t("ent.charges_var"), i.charges_variables,
             t("ent.note_variables")),
            (t("ent.point_equilibre"),
             i.point_mort if i.point_mort is not None else "—",
             t("exp.note_point_mort")),
        ]),
        (t("exp.bloc_tresorerie"), [
            (t("ent.tresorerie"), i.tresorerie, t("exp.note_solde_depart")),
            (t("exp.resultat_mensuel"), i.burn_rate, t("exp.note_burn")),
            (t("ent.autonomie"),
             t("ent.n_mois", n=f"{i.runway_mois:.1f}")
             if i.runway_mois is not None
             else (t("ent.illimitee") if i.resultat_par_mois >= 0 else "—"),
             t("exp.note_runway")),
        ]),
        (t("exp.bloc_clients"), [
            (t("ent.facture"), i.ca_facture, t("ent.note_periode")),
            (t("ent.reste_encaisser"), i.encours_client,
             t("ent.note_non_reglees")),
            (t("ent.dont_retard"), i.retard_client, t("ent.note_echeance")),
            (t("ent.delai_encaiss"),
             t("ent.n_jours", n=f"{i.dso:.0f}") if i.dso is not None else "—",
             t("exp.note_dso")),
            (t("exp.taux_recouvrement"),
             f"{i.taux_recouvrement:.0f} %"
             if i.taux_recouvrement is not None else "—",
             t("exp.note_recouvrement")),
            (t("exp.dependance"), f"{i.dependance_premier_client:.0f} %",
             t("exp.note_dependance")),
        ]),
        (t("exp.bloc_fournisseurs"), [
            (t("ent.facture_fourn"), i.achats_factures, t("ent.note_periode")),
            (t("ent.reste_payer"), i.encours_fournisseur,
             t("ent.note_dettes")),
            (t("ent.delai_paiement"),
             t("ent.n_jours", n=f"{i.dpo:.0f}") if i.dpo is not None else "—",
             t("exp.note_dpo")),
            (t("exp.ecart_financement"),
             t("ent.n_jours", n=f"{i.ecart_de_financement:.0f}")
             if i.ecart_de_financement is not None else "—",
             t("exp.note_ecart")),
        ]),
    ]

    for nom, elements in blocs:
        entete(f, ligne, [nom, "", ""], [34, 18, 46])
        ligne += 1
        for libelle, valeur, note in elements:
            f.cell(row=ligne, column=1, value=libelle).font = normal_f
            c = f.cell(row=ligne, column=2, value=(float(valeur)
                       if isinstance(valeur, (int, float, Decimal)) else valeur))
            c.font = gras_f
            c.alignment = droite
            if isinstance(valeur, (int, float, Decimal)):
                c.number_format = fmt
                c.font = Font(name=POLICE, size=10, bold=True,
                              color=VERT if float(valeur) >= 0 else ROUGE)
            n = f.cell(row=ligne, column=3, value=note)
            n.font = Font(name=POLICE, size=9, italic=True, color="808080")
            for col in (1, 2, 3):
                f.cell(row=ligne, column=col).border = bord
            ligne += 1
        ligne += 1

    entete(f, ligne, [t("exp.ce_que_cela_veut"), "", ""], [34, 18, 46])
    ligne += 1
    for niveau, texte in _phrases(i, t, montant):
        c = f.cell(row=ligne, column=1, value=texte)
        c.font = Font(name=POLICE, size=10,
                      bold=(niveau == "alerte"),
                      color={"alerte": ROUGE, "bon": VERT}.get(niveau, "404040"))
        c.alignment = Alignment(wrap_text=True, vertical="top")
        f.merge_cells(start_row=ligne, start_column=1, end_row=ligne, end_column=3)
        f.row_dimensions[ligne].height = 30
        ligne += 1

    # ---------------- Feuille 2 : clients ----------------
    if i.concentration:
        fc = classeur.create_sheet(t("exp.clients_feuille"))
        fc["A1"] = t("ent.repartition_ca")
        fc["A1"].font = titre_f
        entete(fc, 3, [t("col.client"), t("exp.montant_facture"),
                       t("exp.part_ca")], [38, 20, 14])
        for n, c_ in enumerate(i.concentration, start=4):
            fc.cell(row=n, column=1, value=c_.tiers).font = normal_f
            v = fc.cell(row=n, column=2, value=float(c_.montant))
            v.font = gras_f; v.number_format = fmt; v.alignment = droite
            p = fc.cell(row=n, column=3, value=c_.part)
            p.number_format = "0 %"; p.alignment = droite
            p.font = Font(name=POLICE, size=10, bold=True,
                          color=ROUGE if c_.part > 0.3 else "404040")
            if c_.part > 0.3:
                for col in (1, 2, 3):
                    fc.cell(row=n, column=col).fill = fond_gris

    # ---------------- Feuille 3 : postes ----------------
    if synthese is not None and getattr(synthese, "postes", None):
        fp = classeur.create_sheet(t("exp.postes"))
        fp["A1"] = t("exp.postes_titre")
        fp["A1"].font = titre_f
        entete(fp, 3, [t("exp.poste"), t("col.categorie"), t("exp.nature"),
                       t("col.nombre"), t("col.par_mois")],
               [32, 26, 14, 10, 16])
        for n, p in enumerate(synthese.postes, start=4):
            fp.cell(row=n, column=1, value=p.nom).font = normal_f
            fp.cell(row=n, column=2, value=p.categorie).font = normal_f
            fp.cell(row=n, column=3,
                    value=(t("app.entree") if p.est_une_entree
                           else t("col.charge_fixe") if p.fixe
                           else t("col.variable"))).font = normal_f
            fp.cell(row=n, column=4, value=p.nombre).alignment = droite
            v = fp.cell(row=n, column=5, value=float(p.par_mois))
            v.number_format = fmt; v.alignment = droite
            v.font = Font(name=POLICE, size=10, bold=True,
                          color=VERT if p.est_une_entree else ROUGE)

    # ---------------- Feuille 4 : mouvements ----------------
    if mouvements:
        fm = classeur.create_sheet(t("exp.operations"))
        entete(fm, 1, [t("col.date"), t("exp.libelle"), t("col.montant")],
               [14, 60, 18])
        for n, m in enumerate(sorted(mouvements, key=lambda x: x.jour), start=2):
            fm.cell(row=n, column=1, value=m.jour).number_format = "DD/MM/YYYY"
            fm.cell(row=n, column=2, value=m.libelle).font = normal_f
            v = fm.cell(row=n, column=3, value=float(m.montant))
            v.number_format = fmt; v.alignment = droite
            v.font = Font(name=POLICE, size=10,
                          color=VERT if m.montant > 0 else ROUGE)
        fm.freeze_panes = "A2"
        fm.auto_filter.ref = f"A1:C{len(mouvements) + 1}"

    tampon = io.BytesIO()
    classeur.save(tampon)
    return tampon.getvalue()
